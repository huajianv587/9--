#!/usr/bin/env python3
"""Build the Applied Economics Letters strict v2 DOCX draft."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript/ael_strict_accounting_stress_v2_letter.md"
OUT = ROOT / "manuscript/ael_strict_accounting_stress_v2_submission.docx"
TABLE_DIR = ROOT / "outputs/manuscript_v2_strict"

PRESET = {
    "font": "Calibri",
    "body_size": 11,
    "body_after": 4,
    "body_line_spacing": 1.05,
    "h1_size": 15,
    "h2_size": 12,
    "h1_color": "2E74B5",
    "h2_color": "2E74B5",
    "table_border": "D9D9D9",
    "table_header_fill": "F2F4F7",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), PRESET["table_border"])


def set_table_grid(table, widths: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(col)


def set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def apply_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = PRESET["font"]
    normal.font.size = Pt(PRESET["body_size"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["font"])
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["body_line_spacing"]
    for name, size, color in [
        ("Heading 1", PRESET["h1_size"], PRESET["h1_color"]),
        ("Heading 2", PRESET["h2_size"], PRESET["h2_color"]),
    ]:
        style = doc.styles[name]
        style.font.name = PRESET["font"]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["font"])
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def add_text_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part.strip("`*"))
        run.bold = part.startswith("**") and part.endswith("**")
        run.italic = part.startswith("*") and part.endswith("*") and not run.bold
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(10)


def section_text(markdown: str, heading: str) -> str:
    pattern = rf"^## {re.escape(heading)}\n(.*?)(?=^## |\Z)"
    match = re.search(pattern, markdown, flags=re.M | re.S)
    return match.group(1).strip() if match else ""


def clean_lines(text: str) -> list[str]:
    lines = []
    in_code = False
    for raw in text.splitlines():
        line = raw.strip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            lines.append(f"`{line}`")
            continue
        if not line or line.startswith("### Table"):
            continue
        if line.startswith("Draft status:") or line.startswith("Target use:"):
            continue
        lines.append(line)
    return lines


def parse_md_table(path: Path) -> list[list[str]]:
    rows = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, path: Path, widths: list[float], font_size: float = 7.2) -> None:
    rows = parse_md_table(path)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    set_table_grid(table, widths)
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, PRESET["table_header_fill"])
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.bold = r_idx == 0
            run.font.name = PRESET["font"]
            run.font.size = Pt(font_size)
    doc.add_paragraph()


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    set_portrait(doc.sections[0])
    apply_styles(doc)

    title = markdown.splitlines()[0].lstrip("# ").strip()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = PRESET["font"]
    title_run.font.size = Pt(14)

    for heading in [
        "Abstract",
        "1. Introduction",
        "2. Data and Variables",
        "3. Empirical Specification",
        "4. Results",
        "5. Conclusion",
        "Data Availability",
        "References",
    ]:
        doc.add_heading(heading, level=1 if heading in {"Abstract", "References"} else 2)
        for line in clean_lines(section_text(markdown, heading)):
            add_text_paragraph(doc, line)

    doc.add_page_break()
    table_section = doc.add_section()
    set_landscape(table_section)
    doc.add_heading("Tables", level=1)

    table_specs = [
        ("Table 1. Candidate Samples", "table1_candidate_samples.md", [2.25, 0.9, 0.65, 0.95, 0.85, 0.9, 0.8, 1.0], 7.3),
        ("Table 2. Main and Market-Split Logit Estimates", "table3_main_and_market_split.md", [2.25, 0.75, 0.75, 0.82, 0.9, 0.78, 0.8, 0.72, 0.65], 7.0),
        ("Table 3. Robustness Checks", "table4_robustness.md", [2.9, 0.75, 0.75, 0.9, 0.78, 0.82, 0.7], 7.2),
    ]
    for idx, (caption, filename, widths, size) in enumerate(table_specs):
        if idx:
            doc.add_page_break()
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        cap = para.add_run(caption)
        cap.bold = True
        add_table(doc, TABLE_DIR / filename, widths, size)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
