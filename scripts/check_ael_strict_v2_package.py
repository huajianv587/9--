#!/usr/bin/env python3
"""Audit the Applied Economics Letters strict v2 draft package."""

from __future__ import annotations

from pathlib import Path
import sys
import zipfile

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "manuscript/ael_strict_accounting_stress_v2_letter.md"
DOCX = ROOT / "manuscript/ael_strict_accounting_stress_v2_submission.docx"
PDF = ROOT / "outputs/rendered/ael_strict_v2/ael_strict_accounting_stress_v2_submission.pdf"
PNG_DIR = ROOT / "outputs/rendered/ael_strict_v2"

REQUIRED_MARKERS = [
    "Applied Economics Letters",
    "19,322",
    "7,874",
    "0.711",
    "-7.2 percentage points",
    "firm-clustered",
    "Altman and event labels are treated as robustness or validation evidence",
    "does not support causal language",
    "cannot redistribute raw Capital IQ exports",
]

STALE_MARKERS = [
    "19,402",
    "0.568",
    "0.565",
    "broad financial-stress indicator",
    "Forecast-disagreement measures provide mixed incremental evidence",
    "Chrome-controlled Capital IQ current ASX/SGX/Catalist export",
    "not final submission copy",
]


def fail(message: str) -> None:
    print(f"FAIL: {message}")
    raise SystemExit(1)


def word_count(text: str) -> int:
    return len([token for token in text.replace("\n", " ").split(" ") if token.strip()])


def check_files() -> None:
    missing = [path for path in [MD, DOCX, PDF] if not path.exists()]
    if missing:
        fail(f"missing required artifact(s): {missing}")
    pages = sorted(PNG_DIR.glob("page-*.png"))
    if len(pages) != 6:
        fail(f"expected 6 rendered PNG pages, found {len(pages)}")


def check_markdown() -> None:
    text = MD.read_text(encoding="utf-8")
    wc = word_count(text)
    if wc > 2200:
        fail(f"AEL strict v2 markdown is too long for letter route: {wc} words")
    for marker in REQUIRED_MARKERS:
        if marker not in text:
            fail(f"markdown missing marker: {marker}")
    for marker in STALE_MARKERS:
        if marker in text:
            fail(f"markdown contains stale marker: {marker}")


def check_docx() -> None:
    with zipfile.ZipFile(DOCX) as zf:
        bad = zf.testzip()
        if bad:
            fail(f"DOCX archive integrity failed at {bad}")
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    if len(doc.tables) != 3:
        fail(f"expected 3 DOCX tables, found {len(doc.tables)}")
    if len(doc.sections) != 2:
        fail(f"expected 2 DOCX sections, found {len(doc.sections)}")
    expected_shapes = [(7, 8), (5, 9), (9, 7)]
    shapes = [(len(t.rows), len(t.columns)) for t in doc.tables]
    if shapes != expected_shapes:
        fail(f"unexpected DOCX table shapes: {shapes}")
    for marker in REQUIRED_MARKERS[:7]:
        if marker not in text:
            fail(f"DOCX missing marker: {marker}")
    for marker in STALE_MARKERS:
        if marker in text:
            fail(f"DOCX contains stale marker: {marker}")


def check_pdf() -> None:
    reader = PdfReader(str(PDF))
    if len(reader.pages) != 6:
        fail(f"expected 6 PDF pages, found {len(reader.pages)}")
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    for marker in ["19,322", "0.711", "Table 3. Robustness Checks"]:
        if marker not in text:
            fail(f"PDF missing marker: {marker}")
    for marker in ["19,402", "0.568"]:
        if marker in text:
            fail(f"PDF contains stale marker: {marker}")


def main() -> int:
    check_files()
    check_markdown()
    check_docx()
    check_pdf()
    print("AEL strict v2 package checks passed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
